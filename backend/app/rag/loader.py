import os
from pathlib import Path
from typing import List, Dict, Any, Optional
from pydantic import BaseModel
import pypdf
import docx


class LoadedPage(BaseModel):
    page_number: Optional[int]
    content: str


class LoadedDocument(BaseModel):
    filename: str
    file_type: str
    pages: List[LoadedPage]
    total_characters: int


class DocumentLoader:
    """Multi-format document loader supporting PDF, DOCX, and TXT/Markdown files."""

    @staticmethod
    def load_pdf(file_path: Path) -> List[LoadedPage]:
        """Extract text page-by-page from PDF files."""
        pages = []
        try:
            reader = pypdf.PdfReader(str(file_path))
            for index, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(LoadedPage(page_number=index + 1, content=text.strip()))
        except Exception as e:
            raise ValueError(f"Failed to parse PDF document '{file_path.name}': {str(e)}")
        
        return pages

    @staticmethod
    def load_docx(file_path: Path) -> List[LoadedPage]:
        """Extract text from Word (.docx) documents."""
        try:
            doc = docx.Document(str(file_path))
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            
            combined_text = "\n\n".join(full_text)
            if not combined_text:
                return []
                
            return [LoadedPage(page_number=1, content=combined_text)]
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX document '{file_path.name}': {str(e)}")

    @staticmethod
    def load_txt(file_path: Path) -> List[LoadedPage]:
        """Extract text from plain TXT or Markdown files."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read().strip()
            
            if not content:
                return []
                
            return [LoadedPage(page_number=1, content=content)]
        except Exception as e:
            raise ValueError(f"Failed to parse text document '{file_path.name}': {str(e)}")

    @classmethod
    def load(cls, file_path: Path) -> LoadedDocument:
        """Load document based on file extension."""
        ext = file_path.suffix.lower().lstrip(".")
        filename = file_path.name
        
        if ext == "pdf":
            pages = cls.load_pdf(file_path)
        elif ext in ["docx", "doc"]:
            pages = cls.load_docx(file_path)
        elif ext in ["txt", "md", "markdown", "log", "json", "csv"]:
            pages = cls.load_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: .{ext}")

        if not pages:
            raise ValueError(f"Document '{filename}' appears to be empty or unreadable.")

        total_chars = sum(len(p.content) for p in pages)
        return LoadedDocument(
            filename=filename,
            file_type=ext,
            pages=pages,
            total_characters=total_chars
        )
